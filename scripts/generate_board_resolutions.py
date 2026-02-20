# -*- coding: utf-8 -*-
"""기관투자자별 이사회의사록(주식교환 승인) Word 파일 생성"""
import io, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

if sys.stdout.encoding != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

FLIP_DIR = Path(__file__).parent.parent / "flip_review"

# 7개 한국 기관투자자 정보
INSTITUTIONS = [
    {
        "folder": "01_futureplay",
        "name_kr": "(주)퓨처플레이",
        "representative": "[대표이사명]",
        "shares_kr": "424주",
        "shares_kr_detail": "보통주식 17주, 제1종전환우선주식 407주",
        "shares_us": "22,322.06주",
        "pct": "3.45%",
    },
    {
        "folder": "02_datable",
        "name_kr": "㈜데이터블",
        "representative": "[대표이사명]",
        "shares_kr": "104주",
        "shares_kr_detail": "보통주식 4주, 제1종전환우선주식 100주",
        "shares_us": "5,470.59주",
        "pct": "0.84%",
    },
    {
        "folder": "03_base_dstring",
        "name_kr": "베이스디스트링 벤처투자조합",
        "representative": "[업무집행조합원명]",
        "shares_kr": "311주",
        "shares_kr_detail": "보통주식 13주, 제1종전환우선주식 121주, 제2종전환우선주식 177주",
        "shares_us": "16,366.49주",
        "pct": "2.53%",
    },
    {
        "folder": "04_base_metronome",
        "name_kr": "베이스메트로놈 1호 벤처투자조합",
        "representative": "[업무집행조합원명]",
        "shares_kr": "330주",
        "shares_kr_detail": "보통주식 13주, 제1종전환우선주식 129주, 제2종전환우선주식 188주",
        "shares_us": "17,368.78주",
        "pct": "2.69%",
    },
    {
        "folder": "05_nest_company",
        "name_kr": "네스트컴퍼니 개인투자조합 1호",
        "representative": "[업무집행조합원명]",
        "shares_kr": "66주",
        "shares_kr_detail": "보통주식 3주, 제2종전환우선주식 63주",
        "shares_us": "3,472.46주",
        "pct": "0.53%",
    },
    {
        "folder": "06_cnt_tech_sports",
        "name_kr": "씨엔티테크 제18호 스포츠스타트업 투자조합",
        "representative": "[업무집행조합원명]",
        "shares_kr": "74주",
        "shares_kr_detail": "보통주식 3주, 제2종전환우선주식 71주",
        "shares_us": "3,892.78주",
        "pct": "0.60%",
    },
    {
        "folder": "07_ai_angel_cnt",
        "name_kr": "에이아이엔젤-씨엔티테크 개인투자조합 2호",
        "representative": "[업무집행조합원명]",
        "shares_kr": "74주",
        "shares_kr_detail": "보통주식 3주, 제2종전환우선주식 71주",
        "shares_us": "3,892.78주",
        "pct": "0.60%",
    },
    {
        "folder": "08_primer_sazze",
        "name_kr": "Primer Sazze Fund II, LP",
        "representative": "[General Partner명]",
        "shares_kr": "425주",
        "shares_kr_detail": "보통주식 17주, 제1종전환우선주식 57주, 제2종전환우선주식 351주",
        "shares_us": "22,373.80주",
        "pct": "3.46%",
    },
]

FONT = "맑은 고딕"


def set_font(run, size=11, bold=False):
    """run에 한글 폰트 설정"""
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_paragraph(doc, text, size=11, bold=False, align=None, space_after=6):
    """한글 폰트가 적용된 paragraph 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(space_after)
    return p


def generate_resolution(inst):
    doc = Document()

    # 기본 스타일
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    # 페이지 여백
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- 제목 ---
    add_paragraph(doc, "이 사 회 의 사 록", size=16, bold=True, align="center", space_after=20)

    # --- 기본 정보 ---
    name = inst["name_kr"]

    add_paragraph(doc, f"1. 회사명: {name}", bold=True)
    add_paragraph(doc, f"2. 일  시: 2026년    월    일    시    분")
    add_paragraph(doc, f"3. 장  소: {name} 본점 회의실")
    add_paragraph(doc, f"4. 이사의 수: [  ]명")
    add_paragraph(doc, f"5. 출석 이사: [  ]명", space_after=12)

    add_paragraph(
        doc,
        f"위와 같이 이사 전원이 출석하여 이사회 성립요건을 갖추었으므로, "
        f"의장 {inst['representative']}이(가) 개회를 선언하고 아래 안건을 부의하다.",
        space_after=16,
    )

    # --- 안건 ---
    add_paragraph(doc, "■ 안건: 주식교환 승인의 건", size=12, bold=True, space_after=12)

    add_paragraph(
        doc,
        f"의장이 아래와 같이 {name}이(가) 보유한 ㈜이오스튜디오(EO Studio, 이하 '한국법인') "
        f"주식 전량을 EO Studio Inc.(미국 델라웨어주 설립, 이하 '미국법인')의 주식과 교환하는 건에 대하여 "
        f"설명하고, 이사회의 승인을 요청하다.",
        space_after=12,
    )

    # 교환 상세
    add_paragraph(doc, "1. 교환 대상", bold=True)
    add_paragraph(doc, f"   가. 한국법인 주식: {inst['shares_kr']} ({inst['shares_kr_detail']})")
    add_paragraph(doc, f"   나. 교환비율: 한국법인 주식 1주당 미국법인 주식 약 52.66주")
    add_paragraph(doc, f"   다. 취득할 미국법인 주식: {inst['shares_us']} (지분율 {inst['pct']})", space_after=12)

    add_paragraph(doc, "2. 교환 조건", bold=True)
    add_paragraph(doc, "   가. 한국법인 주당 평가액: 158원 (2025년 11월 30일 기준, DCF 방식)")
    add_paragraph(doc, "   나. 미국법인 주당 평가액: 약 3원 (순자산가치 방식)")
    add_paragraph(doc, "   다. 교환일: Closing Day (미국 변호사가 별도 통지)", space_after=12)

    add_paragraph(doc, "3. 관련 계약서", bold=True)
    add_paragraph(doc, "   가. Amended and Restated Certificate of Incorporation")
    add_paragraph(doc, "   나. Common Stock Exchange Agreement")
    add_paragraph(doc, "   다. Preferred Stock Exchange Agreement")
    add_paragraph(doc, "   라. Investors' Rights Agreement")
    add_paragraph(doc, "   마. Right of First Refusal and Co-Sale Agreement")
    add_paragraph(doc, "   바. Voting Agreement", space_after=12)

    add_paragraph(doc, "4. 위임 사항", bold=True)
    add_paragraph(
        doc,
        f"   {inst['representative']}에게 위 주식교환에 관한 일체의 계약 체결, "
        f"서류 서명 및 외국환신고 등 부수적 절차를 위임한다.",
        space_after=16,
    )

    # --- 의결 ---
    add_paragraph(doc, "■ 의결", size=12, bold=True, space_after=8)
    add_paragraph(
        doc,
        "위 안건에 대하여 출석 이사 전원이 이의 없이 찬성하여 원안대로 가결되었음을 확인하다.",
        space_after=20,
    )

    # --- 날짜 ---
    add_paragraph(doc, "2026년      월      일", align="center", space_after=24)

    # --- 서명란 ---
    add_paragraph(doc, f"{name}", align="center", bold=True, space_after=16)

    add_paragraph(doc, f"의  장:  {inst['representative']}    (인)", align="center", space_after=8)
    add_paragraph(doc, f"이  사:  [이사명]    (인)", align="center", space_after=8)
    add_paragraph(doc, f"이  사:  [이사명]    (인)", align="center", space_after=8)

    # 저장
    filename = f"이사회의사록_주식교환승인_{inst['folder']}.docx"
    output_path = FLIP_DIR / inst["folder"] / filename
    doc.save(str(output_path))
    return output_path


def main():
    for inst in INSTITUTIONS:
        path = generate_resolution(inst)
        print(f"OK  {path.relative_to(FLIP_DIR)}")
    print(f"\nDone! 7개 이사회의사록 초안 생성")


if __name__ == "__main__":
    main()
