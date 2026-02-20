# -*- coding: utf-8 -*-
"""투자자별 체크리스트 Word 파일 생성 (세움 양식 기반)"""
import io, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

if sys.stdout.encoding != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

FLIP_DIR = Path(__file__).parent.parent / "flip_review"
DEADLINE = "2026년 2월 27일(금) 17:00 KST"

CONTRACTS = [
    "EO - Amended and Restated Certificate of Incorporation.docx",
    "EO - Common Stock Exchange Agreement.docx",
    "EO - Preferred Stock Exchange Agreement.docx",
    "EO - Investors' Rights Agreement.docx",
    "EO - Right of First Refusal and Co-Sale Agreement.docx",
    "EO - Voting Agreement.docx",
]

RECIPIENTS = [
    ("00_taeyong_kim",     "김태용 (CEO)",                    "ceo_nonresident",        "Signature Packet (Taeyong Kim).pdf",      None),
    ("01_futureplay",      "퓨처플레이 (FuturePlay)",         "institutional_kr", "Signature Packet (FuturePlay).pdf",              "09. 퓨처플레이-증권취득신고서.docx"),
    ("02_datable",         "데이터블 (Datable)",              "institutional_kr", "Signature Packet (Datable).pdf",                 "10. 데이터블-증권취득신고서.docx"),
    ("03_base_dstring",    "베이스디스트링 (Base Dstring)",     "institutional_kr", "Signature Packet (Base Dstring).pdf",            "13. 베이스디스트링벤처투자조합-증권취득신고서.docx"),
    ("04_base_metronome",  "베이스메트로놈 (Base Metronome)",   "institutional_kr", "Signature Packet (Base Metronome).pdf",          "12. 베이스메트로놈1호벤처투자조합-증권취득신고서.docx"),
    ("05_nest_company",    "네스트컴퍼니 (Nest Company)",      "institutional_kr", "Signature Packet (Next Company).pdf",            "15. 네스트컴퍼니개인투자조합2호-증권취득신고서.docx"),
    ("06_cnt_tech_sports", "씨엔티테크 (CNT Tech Sports)",     "institutional_kr", "Signature Packet (CNT Tech Sports Startup).pdf", "16. 씨엔티테크제18호스포츠스타트업투자조합-증권취득신고서.docx"),
    ("07_ai_angel_cnt",    "에이아이엔젤 (AI Angel CNT)",      "institutional_kr", "Signature Packet (AI Angel - CNT Tech).pdf",     "17. 에이아이엔젤씨엔티테크개인투자조합2호-증권취득신고서.docx"),
    ("08_primer_sazze",    "프라이머사제 (Primer Sazze)",      "institutional_us", "Signature Packet (Primer Sazze).pdf",            None),
    ("09_jae_sik_shin",    "신재식",                          "individual_resident",    "Signature Packet (Jae Sik Shin).pdf",     None),
    ("10_sung_hun_kim",    "김성훈",                          "individual_resident",    "Signature Page (Sung Hun Kim).pdf",       None),
    ("11_wonjun_chang",    "장원준",                          "individual_resident",    "Signature Page (Wonjun Chang).pdf",       None),
    ("12_hyungjoon_yang",  "양형준",                          "individual_resident",    "Signature Page (Hyungjoon Yang).pdf",     None),
    ("13_junyong_suh",     "서준용",                          "individual_junyong",     "Signature Page (Junyong Suh).pdf",        None),
    ("14_seungyoon_lee",   "이승윤",                          "individual_nonresident", "Signature Page (Seungyoon Lee).pdf",      None),
    ("15_sungmoon_cho",    "조성문",                          "individual_nonresident", "Signature Page (Sungmoon Cho).pdf",       None),
    ("16_keeyong_han",     "한기용",                          "individual_nonresident", "Signature Page (Keeyong Han).pdf",        None),
]

TYPE_LABELS = {
    "ceo_nonresident": "CEO / 최대주주 (비거주자, 보유지분율: 약 81%)",
    "institutional_kr": "법인 및 투자조합 (보유지분율: 10% 미만)",
    "institutional_us": "기관투자자 (미국, 비거주자)",
    "individual_resident": "개인 거주자 (보유지분율: 10% 미만)",
    "individual_junyong": "개인 거주자 (보유지분율: 10% 미만)",
    "individual_nonresident": "개인 비거주자",
}

FX_TYPE = {
    "ceo_nonresident": ("거주자의 지정거래외국환은행 해외직접투자신고", "외국환거래규정 제9-5조 제1항"),
    "institutional_kr": ("거주자의 한국은행 증권취득신고", "외국환거래규정 제7-31조 제2항"),
    "individual_resident": ("거주자의 한국은행 증권취득신고", "외국환거래규정 제7-31조 제2항"),
    "individual_junyong": ("거주자의 한국은행 증권취득신고", "외국환거래규정 제7-31조 제2항"),
}

FONT = "맑은 고딕"


# --- Helper functions ---

def set_korean_font(run, font_name=FONT):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(doc, text, size=10, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    set_korean_font(run)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    return p


def set_cell(cell, text, size=9, bold=False, color=None, align=None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    set_korean_font(run)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align == "center":
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_seum_table(doc, rows):
    """세움 양식 테이블: No / 필요서류 항목 / 비고 / 부수 / 원본·사본"""
    table = doc.add_table(rows=len(rows) + 1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    headers = ["No", "필요서류 항목", "비고", "부수", "원본/사본"]
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, size=9, bold=True, align="center")

    # 데이터
    for r_idx, (no, item, note, count, orig) in enumerate(rows):
        row = table.rows[r_idx + 1]
        set_cell(row.cells[0], str(no), size=9, align="center")
        set_cell(row.cells[1], item, size=9)
        set_cell(row.cells[2], note, size=8, color=(80, 80, 80))
        set_cell(row.cells[3], str(count), size=9, align="center")
        set_cell(row.cells[4], orig, size=9, align="center")

    # 열 너비
    widths = [Cm(1.0), Cm(8.5), Cm(4.5), Cm(1.0), Cm(1.5)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    doc.add_paragraph("")
    return table


def add_section_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    set_korean_font(run)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_sub_info(doc, label, value):
    p = doc.add_paragraph()
    run1 = p.add_run(f"  {label}: ")
    run1.font.size = Pt(9)
    run1.font.bold = True
    set_korean_font(run1)
    run2 = p.add_run(value)
    run2.font.size = Pt(9)
    set_korean_font(run2)
    p.paragraph_format.space_after = Pt(2)


# --- 공통 서류 행 ---

def common_rows(start_no):
    return [
        (start_no,     "(공통) 주식교환계약서", "체결 전 최종본", 1, "제공완료"),
        (start_no + 1, "(공통) 이오스튜디오 주주명부", "법인인감 날인본", 1, "제공완료"),
        (start_no + 2, "(공통) 전문평가기관 혹은 공인회계사 등의 주식가치평가서", "날인 최종본", 1, "제공완료"),
    ]


# --- FX 서류 데이터 (신고인 유형별) ---

def fx_rows_ceo():
    """김태용 CEO - 해외직접투자신고"""
    rows = [
        (1,  "거래외국환은행 지정신청서", "은행양식 추후 안내", 1, "원본"),
        (2,  "대고객 설명서 및 확인서", "은행양식 추후 안내", 1, "원본"),
        (3,  "개인(신용)정보 수집·이용·제공 동의서(여신 등)", "은행양식 추후 안내", 1, "원본"),
        (4,  "개인(신용)정보 조회 동의서", "은행양식 추후 안내", 1, "원본"),
        (5,  "위임장", "개인인감 날인", 1, "원본"),
        (6,  "개인인감증명서(전자민원창구용 제출불가)", "신고접수일 기준 3개월 이내 발급본", 1, "원본"),
        (7,  "국세 및 지방세 납세증명서", "신고접수일 기준 유효기간 내", 1, "원본"),
        (8,  "주민등록등본(뒷자리 모두 표기)", "신고접수일 기준 3영업일 이내 발급본", 1, "원본"),
        (9,  "신분증", "-", 1, "PDF"),
    ]
    rows += common_rows(10)
    return rows


def fx_rows_institutional_kr(fx_doc):
    """법인/투자조합 - 증권취득신고"""
    rows = [
        (1, f"증권취득신고서(첨부 파일 참조)", "(법인) 법인인감 날인 / (투자조합) 업무집행조합원 법인인감 날인", 2, "원본"),
        (2, "위임장", "(법인) 법인인감 날인 / (투자조합) 업무집행조합원 법인인감 날인", 1, "원본"),
        (3, "법인인감증명서", "신고접수일 기준 3개월 이내 발급본", 1, "원본"),
        (4, "사업자등록증 혹은 고유번호증", "PDF Scan", 1, "사본"),
        (5, "본건 플립 관련 이사회의사록/조합원총회의사록\n혹은 내부결재문서", "날인 및 서명본", 1, "사본"),
    ]
    rows += common_rows(6)
    return rows


def fx_rows_individual_resident():
    """개인 거주자 - 증권취득신고"""
    rows = [
        (1, "위임장", "개인인감 날인", 1, "원본"),
        (2, "개인인감증명서", "신고접수일 기준 3개월 이내 발급본", 1, "원본"),
        (3, "신분증 사본", "PDF Scan", 1, "사본"),
    ]
    rows += common_rows(4)
    return rows


def fx_rows_individual_junyong():
    """서준용 - 증권취득신고 + 출입국증명서"""
    rows = [
        (1, "위임장", "개인인감 날인", 1, "원본"),
        (2, "개인인감증명서", "신고접수일 기준 3개월 이내 발급본", 1, "원본"),
        (3, "신분증 사본", "PDF Scan", 1, "사본"),
        (4, "출입국증명서(최근 2년)", "-", 1, "원본"),
    ]
    rows += common_rows(5)
    return rows


# --- 메인 생성 함수 ---

def generate_checklist(folder, name, typ, pdf, fx_doc):
    doc = Document()

    # 기본 스타일
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # --- 제목 ---
    add_paragraph(doc, "EO Studio Flip - 체크리스트", size=16, bold=True, align="center")
    doc.add_paragraph("")

    # --- 기본 정보 테이블 ---
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("투자자", name),
        ("유형", TYPE_LABELS[typ]),
        ("제출 기한", DEADLINE),
    ]
    for i, (label, value) in enumerate(info_data):
        set_cell(info_table.rows[i].cells[0], label, size=10, bold=True)
        set_cell(info_table.rows[i].cells[1], value, size=10)
    info_table.rows[0].cells[0].width = Cm(3)
    info_table.rows[0].cells[1].width = Cm(13.5)
    doc.add_paragraph("")

    # --- 설명 ---
    add_paragraph(
        doc,
        "이오스튜디오의 지주회사를 한국에서 미국(EO Studio Inc., Delaware)으로 전환하는 "
        "플립(Flip) 거래입니다. 주주님이 보유한 이오스튜디오 주식이 EO Studio Inc. 주식으로 교환됩니다.",
        size=9, color=(80, 80, 80),
    )
    add_paragraph(
        doc,
        "아래의 사항을 살펴보시고 필요한 서류를 준비하여 기한 내 제출해주시기 바랍니다.",
        size=9, color=(80, 80, 80),
    )
    doc.add_paragraph("")

    # ============================================================
    # Section 1: 첨부 서류
    # ============================================================
    add_section_header(doc, "1. 첨부 서류 (이 폴더에 포함된 파일)")

    included = [(1, pdf, "참고용 (전자서명은 모두싸인으로 별도 진행)", 1, "첨부")]
    for idx, c in enumerate(CONTRACTS):
        included.append((idx + 2, c, "계약서 전문", 1, "첨부"))
    if fx_doc:
        included.append((len(included) + 1, fx_doc, "증권취득신고서", 1, "첨부"))

    add_seum_table(doc, included)

    # ============================================================
    # Section 2: Signature Packet/Page 서명
    # ============================================================
    add_section_header(doc, "2. 서명 (모두싸인 전자서명)")

    sign_rows = [
        (1, "모두싸인으로 전자서명 링크 발송 예정", "이메일 확인", "-", "-"),
        (2, f"{pdf} 전체 서명", "5개 계약서 모두 서명 필요", 1, "-"),
        (3, "날짜 미기입", "Closing Day에 변호사가 기입", "-", "-"),
    ]
    add_seum_table(doc, sign_rows)

    # ============================================================
    # Section 3: 외국환신고 필요서류
    # ============================================================
    if typ in FX_TYPE:
        fx_label, fx_law = FX_TYPE[typ]
        add_section_header(doc, "3. 외국환신고 필요서류")
        add_sub_info(doc, "신고유형", f"{fx_label} ({fx_law})")
        add_sub_info(doc, "취득대상", "EO US 발행 보통주식 또는 전환우선주식")
        doc.add_paragraph("")

        if typ == "ceo_nonresident":
            add_seum_table(doc, fx_rows_ceo())
        elif typ == "institutional_kr":
            add_seum_table(doc, fx_rows_institutional_kr(fx_doc))
        elif typ == "individual_resident":
            add_seum_table(doc, fx_rows_individual_resident())
        elif typ == "individual_junyong":
            add_seum_table(doc, fx_rows_individual_junyong())

    elif typ in ("institutional_us", "individual_nonresident"):
        add_section_header(doc, "3. 외국환신고 필요서류")
        add_paragraph(doc, "비거주자 — 한국 외국환신고 서류 불필요", size=9, color=(120, 120, 120))
        doc.add_paragraph("")

    # ============================================================
    # Section 4: 제출 방법
    # ============================================================
    add_section_header(doc, "4. 제출 방법")

    # 서명: 모두싸인 전자서명 (모든 유형 공통)
    p = doc.add_paragraph("서명: 모두싸인 전자서명 링크로 진행 (별도 우편/스캔 불필요)", style="List Bullet")
    for run in p.runs:
        set_korean_font(run)
        run.font.size = Pt(9)

    # 외국환신고 서류: 원본 우편 (해당 유형만)
    if typ in ("ceo_nonresident", "institutional_kr", "individual_resident", "individual_junyong"):
        p = doc.add_paragraph(
            "외국환신고 서류: 원본 우편 → 서울시 강남구 압구정로28길 9-2, 2층 (EO Studio 안서현)",
            style="List Bullet",
        )
        for run in p.runs:
            set_korean_font(run)
            run.font.size = Pt(9)

    doc.add_paragraph("")

    # ============================================================
    # Section 5: 문의
    # ============================================================
    add_section_header(doc, "5. 문의")
    contacts = [
        "EO Studio 안서현: ash@eoeoeo.net / 010-5382-2553",
        "미국 변호사 Eugene Kang: eugene@venturouscounsel.com",
    ]
    if typ in ("ceo_nonresident", "institutional_kr", "individual_resident", "individual_junyong"):
        contacts.append("한국 변호사 정호석 (법무법인 세움): hoseok.jung@seumlaw.com")
    for c in contacts:
        p = doc.add_paragraph(c, style="List Bullet")
        for run in p.runs:
            set_korean_font(run)
            run.font.size = Pt(9)

    # 저장
    output_path = FLIP_DIR / folder / "CHECKLIST.docx"
    doc.save(str(output_path))
    return output_path


def main():
    for folder, name, typ, pdf, fx_doc in RECIPIENTS:
        # 기존 txt 정리
        for old in ("CHECKLIST.txt", "_INFO.txt"):
            old_path = FLIP_DIR / folder / old
            if old_path.exists():
                old_path.unlink()

        generate_checklist(folder, name, typ, pdf, fx_doc)
        print(f"OK  {folder}/CHECKLIST.docx")

    print(f"\nDone! {len(RECIPIENTS)}개 Word 체크리스트 생성")


if __name__ == "__main__":
    main()
